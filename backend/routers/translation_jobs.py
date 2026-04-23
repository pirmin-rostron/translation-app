import json
import logging
import math
import os
import re
from typing import Optional

from pydantic import BaseModel, Field
import threading
from pathlib import Path
from datetime import date, datetime
from collections import defaultdict

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, Query, Request
from fastapi.responses import FileResponse
from docx import Document as DocxDocument
from sqlalchemy import and_, func
from sqlalchemy.orm import Session

from database import SessionLocal, get_db
from models import (
    Document,
    DocumentBlock,
    DocumentSegment,
    GlossaryTerm,
    Organisation,
    ProcessingStageJob,
    Project,
    SegmentAnnotation,
    TranslationJob,
    TranslationResult,
    UsageEvent,
)
from schemas import (
    DocumentBlockResponse,
    ExportResponse,
    PreviewResponse,
    ExportFileResponse,
    ProcessingStageJobResponse,
    ReviewBlocksPageResponse,
    ReviewSummaryResponse,
    SourceEditRequest,
    SourceEditResponse,
    TranslationJobCreateRequest,
    TranslationProgressResponse,
    ReviewBlockResponse,
    ReviewSegmentResponse,
    SegmentAnnotationResponse,
    SegmentResponse,
    TranslationJobResponse,
    TranslationResultResponse,
    TranslationResultUpdateRequest,
)
from services.auth import get_current_active_user, get_current_org
from services.glossary import glossary_match_in_text, glossary_term_to_match, normalize_optional
from services.usage import AMBIGUITY_RESOLVED, JOB_CREATED, JOB_EXPORTED, TRANSLATION_EDITED, WORDS_TRANSLATED, record_event
from services.webhook import deliver_webhook
from services.translation import SegmentContext, get_translation_provider
from services.translation_memory import (
    TranslationMemoryMatch,
    find_exact_memory_match,
    find_semantic_memory_match,
    store_approved_translation,
)
from tasks import run_translation_pipeline
from limiter import limiter

logger = logging.getLogger(__name__)
router = APIRouter(
    prefix="/translation-jobs",
    tags=["translation-jobs"],
    dependencies=[Depends(get_current_active_user)],
)

TRANSLATION_STAGE = "translation"
AMBIGUITY_STAGE = "ambiguity_detection"
RECONSTRUCTION_STAGE = "reconstruction"
EXPORT_DIR = Path(os.getenv("EXPORT_DIR", "exports"))
SEGMENT_REVIEW_STATES = {"unreviewed", "approved", "edited", "memory_match"}
JOB_STATUS_TRANSLATION_QUEUED = "translation_queued"
JOB_STATUS_TRANSLATING = "translating"
JOB_STATUS_IN_REVIEW = "in_review"
JOB_STATUS_DRAFT_SAVED = "draft_saved"
JOB_STATUS_REVIEW_COMPLETE = "review_complete"
JOB_STATUS_READY_FOR_EXPORT = "ready_for_export"
JOB_STATUS_EXPORTED = "exported"
JOB_STATUS_FAILED = "failed"
TRANSLATION_STYLE_NATURAL = "natural"
TRANSLATION_STYLE_FORMAL = "formal"
TRANSLATION_STYLE_LITERAL = "literal"
TRANSLATION_STYLES = {TRANSLATION_STYLE_NATURAL, TRANSLATION_STYLE_FORMAL, TRANSLATION_STYLE_LITERAL}
JOB_LIFECYCLE_STATUSES = {
    JOB_STATUS_TRANSLATION_QUEUED,
    JOB_STATUS_TRANSLATING,
    JOB_STATUS_IN_REVIEW,
    JOB_STATUS_DRAFT_SAVED,
    JOB_STATUS_REVIEW_COMPLETE,
    JOB_STATUS_READY_FOR_EXPORT,
    JOB_STATUS_EXPORTED,
    JOB_STATUS_FAILED,
}


_RTF_CONTROL_PATTERN = re.compile(r"\\[a-z]+-?\d* ?", re.IGNORECASE)
_RTF_HEX_ESCAPE_PATTERN = re.compile(r"\\'[0-9a-fA-F]{2}")
_EXPORT_FILENAME_PATTERN = re.compile(r"^(?P<prefix>.+)-v(?P<version>\d+)\.(?P<ext>txt|rtf|docx)$")
_RTF_METADATA_FRAGMENT_PATTERN = re.compile(
    r"^(fonttbl|colortbl|stylesheet|info|deff\d+|ansi|ansicpg\d+|viewkind\d+|uc\d+|pard|lang\d+|f\d+|fs\d+|cf\d+|highlight\d+)$",
    re.IGNORECASE,
)
_FONT_NAME_FRAGMENT_PATTERN = re.compile(r"^[A-Za-z0-9 _-]{2,80};$")
SUPPORTED_EXPORT_MODES = {"clean_text", "preserve_formatting"}
SUPPORTED_EXPORT_FORMATS = {"txt", "rtf", "docx"}
SOURCE_EDIT_WARNING = 0.25
SOURCE_EDIT_THRESHOLD = 0.33


def _normalize_translation_style(value: str | None) -> str:
    style = (value or TRANSLATION_STYLE_NATURAL).strip().lower()
    if style not in TRANSLATION_STYLES:
        raise HTTPException(
            status_code=400,
            detail=f"translation_style must be one of: {', '.join(sorted(TRANSLATION_STYLES))}",
        )
    return style


def _queue_stage_job(
    db: Session,
    document_id: int,
    translation_job_id: int,
    stage_name: str,
):
    stage_job = ProcessingStageJob(
        document_id=document_id,
        translation_job_id=translation_job_id,
        stage_name=stage_name,
        status="queued",
    )
    db.add(stage_job)
    return stage_job


def _run_translation_pipeline(translation_job_id: int, user_id: int | None = None):
    db = SessionLocal()
    try:
        job = db.query(TranslationJob).filter(TranslationJob.id == translation_job_id).first()
        if not job:
            return
        from services.events import record_job_event
        record_job_event(db, job.id, "translation_started", "Translation pipeline started")
        db.commit()
        while True:
            stage_job = (
                db.query(ProcessingStageJob)
                .filter(
                    ProcessingStageJob.translation_job_id == translation_job_id,
                    ProcessingStageJob.status == "queued",
                )
                .order_by(ProcessingStageJob.id.asc())
                .first()
            )
            if not stage_job:
                return

            stage_job.status = "running"
            stage_job.attempt_count += 1
            stage_job.started_at = datetime.utcnow()
            stage_job.error_message = None
            db.commit()
            logger.info(
                "Stage start: stage=%s document_id=%d translation_job_id=%d attempt=%d",
                stage_job.stage_name,
                stage_job.document_id,
                translation_job_id,
                stage_job.attempt_count,
            )

            try:
                if stage_job.stage_name == TRANSLATION_STAGE:
                    _execute_translation_stage(db=db, translation_job_id=translation_job_id)
                elif stage_job.stage_name == AMBIGUITY_STAGE:
                    _execute_ambiguity_stage(db=db, translation_job_id=translation_job_id)
                elif stage_job.stage_name == RECONSTRUCTION_STAGE:
                    _execute_reconstruction_stage(db=db, translation_job_id=translation_job_id)
                else:
                    raise ValueError(f"Unknown stage: {stage_job.stage_name}")

                stage_job.status = "succeeded"
                stage_job.finished_at = datetime.utcnow()
                db.commit()
                logger.info(
                    "Stage success: stage=%s document_id=%d translation_job_id=%d",
                    stage_job.stage_name,
                    stage_job.document_id,
                    translation_job_id,
                )
                if stage_job.stage_name == RECONSTRUCTION_STAGE:
                    record_job_event(db, translation_job_id, "translation_complete", "All stages complete, ready for review")
                    db.commit()
                if stage_job.stage_name == TRANSLATION_STAGE:
                    _job = db.query(TranslationJob).filter(TranslationJob.id == translation_job_id).first()
                    word_count = 0
                    if _job:
                        _segs = (
                            db.query(DocumentSegment)
                            .filter(DocumentSegment.document_id == _job.document_id)
                            .all()
                        )
                        word_count = sum(len((s.source_text or "").split()) for s in _segs)
                    record_event(
                        db,
                        WORDS_TRANSLATED,
                        user_id=user_id,
                        job_id=translation_job_id,
                        org_id=_job.org_id if _job else None,
                        meta={"word_count": word_count},
                    )
            except Exception as exc:
                db.rollback()
                job = db.query(TranslationJob).filter(TranslationJob.id == translation_job_id).first()
                if job:
                    job.status = JOB_STATUS_FAILED
                    job.error_message = str(exc)
                failed_stage = db.query(ProcessingStageJob).filter(ProcessingStageJob.id == stage_job.id).first()
                if failed_stage:
                    failed_stage.status = "failed"
                    failed_stage.error_message = str(exc)
                    failed_stage.finished_at = datetime.utcnow()
                record_job_event(db, translation_job_id, "error", str(exc), {"stage": stage_job.stage_name})
                db.commit()
                logger.exception(
                    "Stage failure: stage=%s document_id=%d translation_job_id=%d",
                    stage_job.stage_name,
                    stage_job.document_id,
                    translation_job_id,
                )
                return
    finally:
        db.close()


def _run_translation_pipeline_from_task(translation_job_id: int, user_id: int | None = None, org_id: int | None = None) -> None:
    """Thin wrapper called by the Celery run_translation_pipeline task."""
    _run_translation_pipeline(translation_job_id=translation_job_id, user_id=user_id)


def _get_glossary_candidates(
    db: Session,
    source_language: str,
    target_language: str,
    industry: str | None,
    domain: str | None,
    org_id: int | None = None,
) -> list[GlossaryTerm]:
    normalized_industry = normalize_optional(industry)
    normalized_domain = normalize_optional(domain)
    q = (
        db.query(GlossaryTerm)
        .filter(
            GlossaryTerm.source_language == source_language,
            GlossaryTerm.target_language == target_language,
            GlossaryTerm.industry == normalized_industry,
            GlossaryTerm.domain == normalized_domain,
        )
    )
    if org_id is not None:
        q = q.filter(GlossaryTerm.org_id == org_id)
    return q.order_by(GlossaryTerm.created_at.desc()).all()


def _match_glossary_terms_for_segment(source_text: str, glossary_terms: list[GlossaryTerm]) -> list[dict]:
    matches = []
    for term in glossary_terms:
        if glossary_match_in_text(source_text, term.source_term):
            matches.append(glossary_term_to_match(term))
    return matches


def _build_glossary_matches_payload(glossary_matches: list[dict] | None) -> dict | None:
    if not glossary_matches:
        return None
    return {
        "matches": [
            {
                "source_term": str(match.get("source_term", "")).strip(),
                "target_term": str(match.get("target_term", "")).strip(),
            }
            for match in glossary_matches
            if str(match.get("source_term", "")).strip() and str(match.get("target_term", "")).strip()
        ]
    }


def _normalize_glossary_matches_payload(glossary_matches: object) -> dict | None:
    if glossary_matches is None:
        return None
    if isinstance(glossary_matches, dict):
        matches = glossary_matches.get("matches")
        if isinstance(matches, list):
            return _build_glossary_matches_payload(matches)
        return None
    if isinstance(glossary_matches, list):
        return _build_glossary_matches_payload(glossary_matches)
    return None


def _is_semantic_memory_result(result: TranslationResult) -> bool:
    return bool(getattr(result, "semantic_memory_used", False))


def _is_exact_memory_result(result: TranslationResult) -> bool:
    return bool(getattr(result, "exact_memory_used", False)) or getattr(result, "review_status", "") == "memory_match"


def _has_memory_signal(result: TranslationResult) -> bool:
    return _is_exact_memory_result(result) or _is_semantic_memory_result(result)


def _has_review_signal(result: TranslationResult) -> bool:
    return bool(getattr(result, "ambiguity_detected", False)) or bool(
        getattr(result, "glossary_applied", False)
    ) or _has_memory_signal(result)


def _is_flagged_result(result: TranslationResult) -> bool:
    if _is_acceptable_final_status(getattr(result, "review_status", "")):
        return False
    return _has_review_signal(result)


def _semantic_choice_payload(result: TranslationResult) -> tuple[bool, str | None, float | None, str]:
    details = getattr(result, "semantic_memory_details", None)
    suggested_translation: str | None = None
    similarity_score: float | None = None
    if isinstance(details, dict):
        raw_suggestion = details.get("suggested_translation")
        if isinstance(raw_suggestion, str) and raw_suggestion.strip():
            suggested_translation = raw_suggestion
        raw_similarity = details.get("similarity_score")
        if isinstance(raw_similarity, (float, int)):
            similarity_score = float(raw_similarity)

    current_translation = result.final_translation
    semantic_match_found = bool(_is_semantic_memory_result(result) and suggested_translation)
    return semantic_match_found, suggested_translation, similarity_score, current_translation


def _clean_choice_translation(text: str | None) -> str:
    if not text:
        return ""
    cleaned = text.replace("\\n", " ")
    cleaned = _RTF_HEX_ESCAPE_PATTERN.sub(" ", cleaned)
    cleaned = _RTF_CONTROL_PATTERN.sub(" ", cleaned)
    cleaned = cleaned.replace("{", " ").replace("}", " ")
    return " ".join(cleaned.split()).strip()


def _clean_export_translation(text: str | None) -> str:
    if not text:
        return ""
    cleaned = text.replace("\\n", "\n")
    cleaned = cleaned.replace("\\r\\n", "\n")
    # Preserve paragraph boundaries from RTF-like markers before stripping control words.
    cleaned = re.sub(r"\\par\b", "\n", cleaned, flags=re.IGNORECASE)
    cleaned = _RTF_HEX_ESCAPE_PATTERN.sub(" ", cleaned)
    cleaned = _RTF_CONTROL_PATTERN.sub(" ", cleaned)
    cleaned = cleaned.replace("{", " ").replace("}", " ")
    # Normalize whitespace but keep meaningful line breaks.
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    cleaned = "\n".join(line.strip() for line in cleaned.splitlines())
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _preserve_export_translation(text: str | None) -> str:
    if not text:
        return ""
    preserved = text.replace("\\r\\n", "\n").replace("\\n", "\n")
    preserved = "\n".join(line.rstrip() for line in preserved.splitlines())
    preserved = re.sub(r"\n{3,}", "\n\n", preserved)
    return preserved.strip()


def _clean_review_display_text(text: str | None) -> str:
    if not text:
        return ""
    normalized = text.replace("\\r\\n", "\n").replace("\\n", "\n")
    normalized = re.sub(r"\\par\b", "\n", normalized, flags=re.IGNORECASE)
    normalized = _RTF_HEX_ESCAPE_PATTERN.sub(" ", normalized)
    normalized = _RTF_CONTROL_PATTERN.sub(" ", normalized)
    normalized = normalized.replace("{", " ").replace("}", " ")
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    lines = [line.strip() for line in normalized.splitlines() if line.strip()]
    filtered_lines: list[str] = []
    for line in lines:
        if _RTF_METADATA_FRAGMENT_PATTERN.match(line):
            continue
        if _FONT_NAME_FRAGMENT_PATTERN.match(line):
            continue
        filtered_lines.append(line)
    if not filtered_lines:
        return ""
    return "\n".join(filtered_lines).strip()


def _normalize_export_mode(export_mode: str) -> str:
    normalized = (export_mode or "").strip().lower()
    if normalized not in SUPPORTED_EXPORT_MODES:
        raise HTTPException(
            status_code=400,
            detail=f"export_mode must be one of: {', '.join(sorted(SUPPORTED_EXPORT_MODES))}",
        )
    return normalized


def _ambiguity_choice_payload(result: TranslationResult) -> tuple[bool, str | None, list[dict[str, str]]]:
    details = getattr(result, "ambiguity_details", None)
    source_phrase: str | None = None
    options: list[dict[str, str]] = []
    seen_translations: set[str] = set()

    if isinstance(details, dict):
        raw_source_phrase = details.get("source_span")
        if isinstance(raw_source_phrase, str) and raw_source_phrase.strip():
            source_phrase = raw_source_phrase.strip()
        raw_alternatives = details.get("alternatives")
        if isinstance(raw_alternatives, list):
            for idx, option in enumerate(raw_alternatives):
                if not isinstance(option, dict):
                    continue
                translation = _clean_choice_translation(str(option.get("translation", "")))
                if not translation:
                    continue
                meaning = str(option.get("meaning", "")).strip() or f"Possible meaning {idx + 1}"
                normalized_translation = translation.casefold()
                if normalized_translation in seen_translations:
                    continue
                seen_translations.add(normalized_translation)
                options.append({"meaning": meaning, "translation": translation})

    ambiguity_choice_found = bool(getattr(result, "ambiguity_detected", False) and options)
    return ambiguity_choice_found, source_phrase, options


def _serialize_translation_result(result: TranslationResult, segment: DocumentSegment | None) -> TranslationResultResponse:
    seg_resp = SegmentResponse.model_validate(segment) if segment else None
    semantic_match_found, suggested_translation, similarity_score, current_translation = _semantic_choice_payload(result)
    ambiguity_choice_found, ambiguity_source_phrase, ambiguity_options = _ambiguity_choice_payload(result)
    return TranslationResultResponse(
        id=result.id,
        job_id=result.job_id,
        segment_id=result.segment_id,
        primary_translation=result.primary_translation,
        final_translation=result.final_translation,
        confidence_score=result.confidence_score,
        review_status=_normalize_review_status(result.review_status),
        exact_memory_used=_is_exact_memory_result(result),
        semantic_memory_used=_is_semantic_memory_result(result),
        semantic_memory_details=getattr(result, "semantic_memory_details", None),
        semantic_match_found=semantic_match_found,
        suggested_translation=suggested_translation,
        similarity_score=similarity_score,
        current_translation=current_translation,
        ambiguity_choice_found=ambiguity_choice_found,
        ambiguity_source_phrase=ambiguity_source_phrase,
        ambiguity_options=ambiguity_options,
        ambiguity_detected=getattr(result, "ambiguity_detected", False) or False,
        ambiguity_details=getattr(result, "ambiguity_details", None),
        glossary_applied=getattr(result, "glossary_applied", False) or False,
        glossary_matches=_normalize_glossary_matches_payload(getattr(result, "glossary_matches", None)),
        created_at=result.created_at,
        segment=seg_resp,
    )


def _find_span(text: str, needle: str) -> tuple[int, int] | None:
    text_value = (text or "").strip()
    needle_value = (needle or "").strip()
    if not text_value or not needle_value:
        return None
    start = text_value.lower().find(needle_value.lower())
    if start == -1:
        return None
    return start, start + len(needle_value)


def _add_annotation(
    annotations: list[SegmentAnnotation],
    segment_id: int,
    translation_job_id: int,
    annotation_type: str,
    source_text: str,
    source_span_text: str,
    target_text: str,
    target_span_text: str | None,
    metadata_json: dict | None = None,
):
    source_span = _find_span(source_text, source_span_text)
    if not source_span:
        return

    target_span = _find_span(target_text, target_span_text) if target_span_text else None
    annotations.append(
        SegmentAnnotation(
            segment_id=segment_id,
            translation_job_id=translation_job_id,
            annotation_type=annotation_type,
            source_span_text=source_span_text,
            source_start=source_span[0],
            source_end=source_span[1],
            target_span_text=target_span_text,
            target_start=target_span[0] if target_span else None,
            target_end=target_span[1] if target_span else None,
            metadata_json=metadata_json,
        )
    )


def _build_segment_annotations(
    segment: DocumentSegment,
    result: TranslationResult,
) -> list[SegmentAnnotation]:
    annotations: list[SegmentAnnotation] = []
    source_text = segment.source_text
    target_text = result.final_translation
    glossary_matches = _normalize_glossary_matches_payload(getattr(result, "glossary_matches", None))

    if glossary_matches:
        for match in glossary_matches.get("matches", []):
            source_term = str(match.get("source_term", "")).strip()
            target_term = str(match.get("target_term", "")).strip()
            if not source_term:
                continue
            _add_annotation(
                annotations=annotations,
                segment_id=segment.id,
                translation_job_id=result.job_id,
                annotation_type="glossary",
                source_text=source_text,
                source_span_text=source_term,
                target_text=target_text,
                target_span_text=target_term or None,
                metadata_json={
                    "source_term": source_term,
                    "target_term": target_term or None,
                },
            )

    if getattr(result, "ambiguity_detected", False) and getattr(result, "ambiguity_details", None):
        ambiguity_details = result.ambiguity_details or {}
        source_span_text = str(ambiguity_details.get("source_span", "")).strip()
        target_span_text = None
        for alternative in ambiguity_details.get("alternatives", []):
            candidate = str(alternative.get("translation", "")).strip()
            if candidate and _find_span(target_text, candidate):
                target_span_text = candidate
                break
        if source_span_text:
            _add_annotation(
                annotations=annotations,
                segment_id=segment.id,
                translation_job_id=result.job_id,
                annotation_type="ambiguity",
                source_text=source_text,
                source_span_text=source_span_text,
                target_text=target_text,
                target_span_text=target_span_text,
                metadata_json=ambiguity_details,
            )

    if _is_exact_memory_result(result):
        _add_annotation(
            annotations=annotations,
            segment_id=segment.id,
            translation_job_id=result.job_id,
            annotation_type="exact_memory",
            source_text=source_text,
            source_span_text=source_text,
            target_text=target_text,
            target_span_text=target_text,
            metadata_json={"label": "Suggested from previously approved translation"},
        )

    if _is_semantic_memory_result(result):
        _add_annotation(
            annotations=annotations,
            segment_id=segment.id,
            translation_job_id=result.job_id,
            annotation_type="semantic_memory",
            source_text=source_text,
            source_span_text=source_text,
            target_text=target_text,
            target_span_text=target_text,
            metadata_json={
                "label": "Similar approved translation reused",
                "confidence_score": result.confidence_score,
            },
        )

    return annotations


def _replace_segment_annotations(db: Session, segment: DocumentSegment, result: TranslationResult):
    db.query(SegmentAnnotation).filter(
        SegmentAnnotation.segment_id == segment.id,
        SegmentAnnotation.translation_job_id == result.job_id,
    ).delete()
    annotations = _build_segment_annotations(segment, result)
    if annotations:
        db.add_all(annotations)


def _compose_block_translation(block_type: str, parts: list[str]) -> str | None:
    cleaned = [part.strip() for part in parts if part and part.strip()]
    if not cleaned:
        return None
    separator = "\n" if block_type == "bullet_item" else " "
    return separator.join(cleaned)


def _compose_review_block_translation(parts: list[str]) -> str | None:
    cleaned = [part.strip() for part in parts if part and part.strip()]
    if not cleaned:
        return None
    # Preserve line-level structure in review payload to avoid collapsing heading/sentence boundaries.
    return "\n".join(cleaned)


def _build_review_block_translated_representations(
    block: DocumentBlock,
    block_segments: list[DocumentSegment],
    result_by_segment_id: dict[int, TranslationResult],
) -> tuple[str, str | None, str]:
    """Build canonical translated block representations for review and preview.

    - translated_text_raw: markup-preserving view, intentionally composed from segment
      primary translations (fallback to final segment text when needed).
    - translated_text_display: clean readable view from reviewed final translations.
    - translated_text_final: canonical reviewed text used by review/export logic.
    """
    final_parts: list[str] = []
    raw_parts: list[str] = []
    for segment in block_segments:
        result = result_by_segment_id.get(segment.id)
        if not result:
            continue
        final_text = (result.final_translation or "").strip()
        if final_text:
            final_parts.append(final_text)
        raw_text = (result.primary_translation or "").strip()
        if raw_text:
            raw_parts.append(raw_text)
        elif final_text:
            raw_parts.append(final_text)

    translated_text_final = _compose_review_block_translation(final_parts) or (block.text_translated or "").strip()
    translated_text_raw = _compose_review_block_translation(raw_parts) or translated_text_final
    translated_text_display = _clean_review_display_text(translated_text_final) if translated_text_final else None
    return translated_text_raw, translated_text_display, translated_text_final


def _refresh_document_block_translation(db: Session, block_id: int, job_id: int):
    block = db.query(DocumentBlock).filter(DocumentBlock.id == block_id).first()
    if not block:
        return

    segments = (
        db.query(DocumentSegment)
        .filter(DocumentSegment.block_id == block_id)
        .order_by(DocumentSegment.segment_index)
        .all()
    )
    translated_parts = []
    for segment in segments:
        result = (
            db.query(TranslationResult)
            .filter(TranslationResult.job_id == job_id, TranslationResult.segment_id == segment.id)
            .first()
        )
        if result:
            raw = (result.final_translation or "").strip()
            if raw:
                translated_parts.append(raw)
    block.text_translated = _compose_review_block_translation(translated_parts)


def _serialize_review_segment(
    segment: DocumentSegment,
    result: TranslationResult,
    annotations: list[SegmentAnnotation],
) -> ReviewSegmentResponse:
    semantic_match_found, suggested_translation, similarity_score, current_translation = _semantic_choice_payload(result)
    ambiguity_choice_found, ambiguity_source_phrase, ambiguity_options = _ambiguity_choice_payload(result)
    return ReviewSegmentResponse(
        id=result.id,
        segment_id=segment.id,
        block_id=segment.block_id,
        segment_index=segment.segment_index,
        segment_type=segment.segment_type,
        source_text=segment.source_text,
        primary_translation=result.primary_translation,
        final_translation=result.final_translation,
        confidence_score=result.confidence_score,
        review_status=_normalize_review_status(result.review_status),
        exact_memory_used=_is_exact_memory_result(result),
        semantic_memory_used=_is_semantic_memory_result(result),
        semantic_memory_details=getattr(result, "semantic_memory_details", None),
        semantic_match_found=semantic_match_found,
        suggested_translation=suggested_translation,
        similarity_score=similarity_score,
        current_translation=current_translation,
        ambiguity_choice_found=ambiguity_choice_found,
        ambiguity_source_phrase=ambiguity_source_phrase,
        ambiguity_options=ambiguity_options,
        ambiguity_detected=getattr(result, "ambiguity_detected", False) or False,
        ambiguity_details=getattr(result, "ambiguity_details", None),
        glossary_applied=getattr(result, "glossary_applied", False) or False,
        glossary_matches=_normalize_glossary_matches_payload(getattr(result, "glossary_matches", None)),
        annotations=[SegmentAnnotationResponse.model_validate(annotation) for annotation in annotations],
    )


def _is_acceptable_final_status(review_status: str) -> bool:
    return _normalize_review_status(review_status) in {"approved", "edited", "memory_match"}


def _normalize_review_status(review_status: str) -> str:
    # Backward compatibility for historical saved values.
    if review_status == "reviewed":
        return "edited"
    if review_status == "semantic_memory_match":
        return "memory_match"
    return review_status


def _calculate_review_summary(db: Session, job: TranslationJob) -> ReviewSummaryResponse:
    results = db.query(TranslationResult).filter(TranslationResult.job_id == job.id).all()
    total_segments = len(results)
    approved_segments = sum(1 for result in results if _normalize_review_status(result.review_status) == "approved")
    edited_segments = sum(1 for result in results if _normalize_review_status(result.review_status) == "edited")
    unresolved_count = sum(1 for result in results if not _is_acceptable_final_status(result.review_status))
    unresolved_ambiguities = sum(
        1
        for result in results
        if bool(result.ambiguity_detected) and not _is_acceptable_final_status(result.review_status)
    )
    unresolved_semantic_reviews = sum(
        1
        for result in results
        if bool(result.semantic_memory_used) and not _is_acceptable_final_status(result.review_status)
    )
    # Count safe unresolved BLOCKS (not segments) — group results by block
    seg_to_block: dict[int, int] = {}
    for seg in db.query(DocumentSegment).filter(
        DocumentSegment.id.in_([r.segment_id for r in results])
    ).all():
        if seg.block_id is not None:
            seg_to_block[seg.id] = seg.block_id
    block_results: dict[int, list] = {}
    for result in results:
        bid = seg_to_block.get(result.segment_id, result.segment_id)
        block_results.setdefault(bid, []).append(result)
    safe_unresolved_segments = 0
    for block_segs in block_results.values():
        has_unresolved = any(not _is_acceptable_final_status(r.review_status) for r in block_segs)
        if not has_unresolved:
            continue
        all_safe = all(
            not bool(r.ambiguity_detected)
            and not bool(r.semantic_memory_used)
            and bool((r.final_translation or "").strip())
            for r in block_segs
            if not _is_acceptable_final_status(r.review_status)
        )
        if all_safe:
            safe_unresolved_segments += 1
    review_complete = (
        total_segments > 0
        and unresolved_count == 0
        and unresolved_ambiguities == 0
        and unresolved_semantic_reviews == 0
    )
    can_mark_ready_for_export = review_complete
    overall_status = job.status if job.status in JOB_LIFECYCLE_STATUSES else JOB_STATUS_IN_REVIEW
    return ReviewSummaryResponse(
        job_id=job.id,
        total_segments=total_segments,
        approved_segments=approved_segments,
        edited_segments=edited_segments,
        safe_unresolved_segments=safe_unresolved_segments,
        review_complete=review_complete,
        unresolved_count=unresolved_count,
        unresolved_ambiguities=unresolved_ambiguities,
        unresolved_semantic_reviews=unresolved_semantic_reviews,
        unresolved_segments=unresolved_count,
        ambiguity_count=unresolved_ambiguities,
        semantic_memory_review_count=unresolved_semantic_reviews,
        overall_status=overall_status,
        last_saved_at=job.last_saved_at,
        can_mark_ready_for_export=can_mark_ready_for_export,
    )


def _collect_export_blocks(db: Session, job: TranslationJob, export_mode: str) -> list[tuple[str, str]]:
    blocks = (
        db.query(DocumentBlock)
        .filter(DocumentBlock.document_id == job.document_id)
        .order_by(DocumentBlock.block_index)
        .all()
    )
    segments = (
        db.query(DocumentSegment)
        .filter(DocumentSegment.document_id == job.document_id)
        .order_by(DocumentSegment.segment_index)
        .all()
    )
    results = db.query(TranslationResult).filter(TranslationResult.job_id == job.id).all()
    result_by_segment_id = {result.segment_id: result for result in results}
    segments_by_block_id: dict[int | None, list[DocumentSegment]] = defaultdict(list)
    for segment in segments:
        segments_by_block_id[segment.block_id].append(segment)

    output_blocks: list[tuple[str, str]] = []
    for block in blocks:
        block_segments = sorted(
            segments_by_block_id.get(block.id, []),
            key=lambda segment: segment.segment_index,
        )
        translated_parts: list[str] = []
        for segment in block_segments:
            result = result_by_segment_id.get(segment.id)
            if not result:
                continue
            cleaned_part = (
                _clean_export_translation(result.final_translation)
                if export_mode == "clean_text"
                else _preserve_export_translation(result.final_translation)
            )
            if not cleaned_part:
                continue
            translated_parts.append(cleaned_part)
        translated_text = _compose_block_translation(block.block_type, translated_parts) or ""
        translated_text = translated_text.strip()
        if not translated_text:
            continue
        output_blocks.append((block.block_type or "paragraph", translated_text))
    return output_blocks


def _collect_preview_blocks(db: Session, job: TranslationJob) -> list[tuple[str, str, str]]:
    blocks = (
        db.query(DocumentBlock)
        .filter(DocumentBlock.document_id == job.document_id)
        .order_by(DocumentBlock.block_index)
        .all()
    )
    segments = (
        db.query(DocumentSegment)
        .filter(DocumentSegment.document_id == job.document_id)
        .order_by(DocumentSegment.segment_index)
        .all()
    )
    results = db.query(TranslationResult).filter(TranslationResult.job_id == job.id).all()
    result_by_segment_id = {result.segment_id: result for result in results}
    segments_by_block_id: dict[int | None, list[DocumentSegment]] = defaultdict(list)
    for segment in segments:
        segments_by_block_id[segment.block_id].append(segment)

    output_blocks: list[tuple[str, str, str]] = []
    for block in blocks:
        block_segments = sorted(
            segments_by_block_id.get(block.id, []),
            key=lambda segment: segment.segment_index,
        )
        translated_text_raw, translated_text_display, _ = _build_review_block_translated_representations(
            block=block,
            block_segments=block_segments,
            result_by_segment_id=result_by_segment_id,
        )
        if not (translated_text_raw or translated_text_display):
            continue
        output_blocks.append(
            (
                block.block_type or "paragraph",
                translated_text_raw.strip(),
                (translated_text_display or "").strip(),
            )
        )
    return output_blocks


def _build_export_text(db: Session, job: TranslationJob, export_mode: str) -> str:
    output_blocks = _collect_export_blocks(db, job, export_mode)
    output_lines: list[str] = []
    for block_type, translated_text in output_blocks:
        output_lines.append(f"- {translated_text}" if block_type == "bullet_item" else translated_text)
        output_lines.append("")
    return "\n".join(output_lines).strip() + "\n"


def _render_preview_document_text(output_blocks: list[tuple[str, str, str]], view_mode: str) -> str:
    sections: list[str] = []
    for block_type, translated_raw, translated_display in output_blocks:
        text_value = translated_raw if view_mode == "raw" else translated_display
        text_value = text_value.strip()
        if not text_value:
            continue
        sections.append(f"- {text_value}" if block_type == "bullet_item" else text_value)
    return "\n\n".join(sections).strip()


def _build_export_rtf(db: Session, job: TranslationJob, export_mode: str) -> str:
    def _rtf_escape(value: str) -> str:
        escaped = (
            value.replace("\\", "\\\\")
            .replace("{", "\\{")
            .replace("}", "\\}")
            .replace("\n", "\\line ")
        )
        return escaped

    output_blocks = _collect_export_blocks(db, job, export_mode)
    rtf_parts = ["{\\rtf1\\ansi\\deff0", "\\viewkind4\\uc1\\pard"]
    for block_type, translated_text in output_blocks:
        escaped = _rtf_escape(translated_text)
        if export_mode == "preserve_formatting":
            if block_type == "heading":
                rtf_parts.append(f"\\b {escaped}\\b0\\par")
            elif block_type == "bullet_item":
                rtf_parts.append(f"\\tab \\bullet\\tab {escaped}\\par")
            else:
                rtf_parts.append(f"{escaped}\\par")
        else:
            plain = _rtf_escape(_clean_export_translation(translated_text))
            rtf_parts.append(f"{plain}\\par")
    rtf_parts.append("}")
    return "\n".join(rtf_parts) + "\n"


def _build_export_docx(db: Session, job: TranslationJob, export_mode: str, output_path: Path) -> None:
    output_blocks = _collect_export_blocks(db, job, export_mode)
    doc = DocxDocument()
    for block_type, translated_text in output_blocks:
        text_value = translated_text if export_mode == "preserve_formatting" else _clean_export_translation(translated_text)
        if not text_value:
            continue
        if export_mode == "preserve_formatting":
            if block_type == "heading":
                doc.add_heading(text_value, level=2)
            elif block_type == "bullet_item":
                doc.add_paragraph(text_value, style="List Bullet")
            else:
                doc.add_paragraph(text_value)
        else:
            doc.add_paragraph(text_value)
    doc.save(str(output_path))


def _export_filename_prefix(doc: Document, job: TranslationJob) -> str:
    safe_stem = Path(doc.filename).stem or f"document-{doc.id}"
    return f"{safe_stem}-job-{job.id}"


def _export_metadata_path(filename: str) -> Path:
    return EXPORT_DIR / f"{filename}.meta.json"


def _write_export_metadata(filename: str, generated_at: datetime, export_format: str, export_mode: str) -> None:
    metadata = {
        "generated_at": generated_at.isoformat(),
        "export_format": export_format,
        "export_mode": export_mode,
    }
    _export_metadata_path(filename).write_text(json.dumps(metadata), encoding="utf-8")


def _read_export_metadata(filename: str) -> dict[str, str] | None:
    path = _export_metadata_path(filename)
    if not path.exists() or not path.is_file():
        return None
    try:
        raw = path.read_text(encoding="utf-8")
        loaded = json.loads(raw)
        if isinstance(loaded, dict):
            return {
                "generated_at": str(loaded.get("generated_at", "")).strip(),
                "export_format": str(loaded.get("export_format", "")).strip(),
                "export_mode": str(loaded.get("export_mode", "")).strip(),
            }
    except Exception:
        logger.warning("Failed to read export metadata for %s", filename, exc_info=True)
    return None


def _list_export_files(job: TranslationJob, doc: Document) -> list[ExportFileResponse]:
    prefix = _export_filename_prefix(doc, job)
    files = sorted(EXPORT_DIR.glob(f"{prefix}-v*.*"))
    parsed: list[ExportFileResponse] = []
    for path in files:
        match = _EXPORT_FILENAME_PATTERN.match(path.name)
        if not match:
            continue
        if match.group("prefix") != prefix:
            continue
        version = int(match.group("version"))
        metadata = _read_export_metadata(path.name)
        generated_at = datetime.utcfromtimestamp(path.stat().st_mtime)
        export_format: str | None = match.group("ext")
        export_mode: str | None = None
        if metadata:
            generated_at_raw = metadata.get("generated_at")
            if generated_at_raw:
                try:
                    generated_at = datetime.fromisoformat(generated_at_raw)
                except ValueError:
                    pass
            export_format_raw = metadata.get("export_format")
            export_mode_raw = metadata.get("export_mode")
            if export_format_raw:
                export_format = export_format_raw
            if export_mode_raw:
                export_mode = export_mode_raw
        parsed.append(
            ExportFileResponse(
                filename=path.name,
                download_url=f"/translation-jobs/{job.id}/exports/{path.name}",
                generated_at=generated_at,
                version=version,
                export_format=export_format,
                export_mode=export_mode,
            )
        )

    parsed.sort(key=lambda item: item.version, reverse=True)
    if parsed:
        parsed[0].latest = True
    return parsed


def _next_export_version(job: TranslationJob, doc: Document) -> int:
    existing = _list_export_files(job, doc)
    if not existing:
        return 1
    return max(item.version for item in existing) + 1


def _estimate_translation_eta_seconds(job: TranslationJob, total_segments: int, completed_segments: int) -> int | None:
    if total_segments <= 0 or completed_segments <= 0 or not job.progress_started_at:
        return None
    elapsed = (datetime.utcnow() - job.progress_started_at).total_seconds()
    if elapsed <= 0:
        return None
    rate = completed_segments / elapsed
    if rate <= 0:
        return None
    remaining = total_segments - completed_segments
    if remaining <= 0:
        return 0
    return max(int(remaining / rate), 1)


def _calculate_translation_progress(db: Session, job: TranslationJob) -> TranslationProgressResponse:
    stage_jobs = (
        db.query(ProcessingStageJob)
        .filter(ProcessingStageJob.translation_job_id == job.id)
        .order_by(ProcessingStageJob.id.asc())
        .all()
    )
    translation_stage = next((stage for stage in stage_jobs if stage.stage_name == TRANSLATION_STAGE), None)
    ambiguity_stage = next((stage for stage in stage_jobs if stage.stage_name == AMBIGUITY_STAGE), None)
    reconstruction_stage = next((stage for stage in stage_jobs if stage.stage_name == RECONSTRUCTION_STAGE), None)

    total_segments = job.progress_total_segments or (
        db.query(DocumentSegment).filter(DocumentSegment.document_id == job.document_id).count()
    )
    completed_segments = max(job.progress_completed_segments or 0, 0)
    eta_seconds: int | None = None
    stage_label = "Translation queued"
    percentage = 0.0

    # Block counts: number of blocks whose every segment has a TranslationResult for this job.
    _block_subq = (
        db.query(
            DocumentSegment.block_id,
            func.count(DocumentSegment.id).label("seg_count"),
            func.count(TranslationResult.id).label("result_count"),
        )
        .outerjoin(
            TranslationResult,
            and_(
                TranslationResult.segment_id == DocumentSegment.id,
                TranslationResult.job_id == job.id,
            ),
        )
        .filter(
            DocumentSegment.document_id == job.document_id,
            DocumentSegment.block_id.isnot(None),
        )
        .group_by(DocumentSegment.block_id)
        .subquery()
    )
    blocks_total: int = db.query(func.count()).select_from(_block_subq).scalar() or 0
    blocks_completed: int = (
        db.query(func.count())
        .select_from(_block_subq)
        .filter(_block_subq.c.seg_count == _block_subq.c.result_count)
        .scalar()
    ) or 0

    completed_statuses = {
        JOB_STATUS_IN_REVIEW,
        JOB_STATUS_DRAFT_SAVED,
        JOB_STATUS_REVIEW_COMPLETE,
        JOB_STATUS_READY_FOR_EXPORT,
        JOB_STATUS_EXPORTED,
    }
    if job.status in completed_statuses:
        return TranslationProgressResponse(
            job_id=job.id,
            stage_label="Translation complete",
            total_segments=total_segments,
            completed_segments=total_segments,
            percentage=100.0,
            eta_seconds=0,
            is_complete=True,
            blocks_completed=blocks_total,
            blocks_total=blocks_total,
        )

    if translation_stage and translation_stage.status in {"running", "queued"}:
        if translation_stage.status == "queued":
            stage_label = "Preparing translation"
            percentage = 5.0
            eta_seconds = None
        else:
            stage_label = f"Translating {completed_segments} of {total_segments} segments"
            segment_pct = (completed_segments / total_segments * 100.0) if total_segments > 0 else 0.0
            percentage = min(segment_pct * 0.85, 85.0)
            eta_seconds = _estimate_translation_eta_seconds(job, total_segments, completed_segments)
    elif ambiguity_stage and ambiguity_stage.status in {"running", "queued"}:
        stage_label = "Detecting ambiguities"
        percentage = 92.0 if ambiguity_stage.status == "running" else 88.0
        eta_seconds = 6 if ambiguity_stage.status == "queued" else 3
    elif reconstruction_stage and reconstruction_stage.status in {"running", "queued"}:
        stage_label = "Preparing review document"
        percentage = 97.0 if reconstruction_stage.status == "running" else 95.0
        eta_seconds = 2
    else:
        if job.status == JOB_STATUS_TRANSLATION_QUEUED:
            stage_label = "Translation queued"
            percentage = 0.0
        elif job.status == JOB_STATUS_TRANSLATING:
            segment_pct = (completed_segments / total_segments * 100.0) if total_segments > 0 else 0.0
            stage_label = f"Translating {completed_segments} of {total_segments} segments"
            percentage = min(segment_pct, 99.0)
            eta_seconds = _estimate_translation_eta_seconds(job, total_segments, completed_segments)
        elif job.status == JOB_STATUS_FAILED:
            stage_label = "Translation failed"
            percentage = 100.0
            eta_seconds = None

    return TranslationProgressResponse(
        job_id=job.id,
        stage_label=stage_label,
        total_segments=total_segments,
        completed_segments=min(completed_segments, total_segments) if total_segments > 0 else completed_segments,
        percentage=round(max(min(percentage, 100.0), 0.0), 1),
        eta_seconds=eta_seconds,
        is_complete=False,
        blocks_completed=blocks_completed,
        blocks_total=blocks_total,
    )


def _execute_translation_stage(db: Session, translation_job_id: int):
    job = db.query(TranslationJob).filter(TranslationJob.id == translation_job_id).first()
    if not job:
        raise ValueError("Translation job not found")
    doc = db.query(Document).filter(Document.id == job.document_id).first()
    if not doc:
        raise ValueError("Document not found")

    existing_count = (
        db.query(TranslationResult)
        .filter(TranslationResult.job_id == translation_job_id)
        .count()
    )
    if existing_count > 0:
        logger.warning(
            "Translation stage called on job_id=%d which already has %d results — skipping to prevent duplicates",
            translation_job_id,
            existing_count,
        )
        return

    job.status = JOB_STATUS_TRANSLATING
    job.error_message = None
    db.commit()

    segments = (
        db.query(DocumentSegment)
        .filter(DocumentSegment.document_id == doc.id)
        .order_by(DocumentSegment.segment_index)
        .all()
    )
    if not segments:
        raise ValueError("Document has no segmented content to translate")

    provider, provider_name = get_translation_provider()
    source_lang = doc.source_language or "unknown"
    translation_style = (job.translation_style or TRANSLATION_STYLE_NATURAL).strip().lower()
    if translation_style not in TRANSLATION_STYLES:
        translation_style = TRANSLATION_STYLE_NATURAL
    glossary_candidates = _get_glossary_candidates(
        db=db,
        source_language=source_lang,
        target_language=doc.target_language,
        industry=doc.industry,
        domain=doc.domain,
        org_id=job.org_id,
    )
    logger.info(
        "Translation stage start: provider=%s style=%s segments=%d document_id=%d job_id=%d",
        provider_name,
        translation_style,
        len(segments),
        doc.id,
        translation_job_id,
    )

    db.query(TranslationResult).filter(TranslationResult.job_id == translation_job_id).delete(
        synchronize_session=False
    )
    db.commit()

    total_segments = len(segments)
    job.progress_total_segments = total_segments
    job.progress_completed_segments = 0
    job.progress_started_at = datetime.utcnow()
    db.commit()

    start_time = datetime.utcnow()
    logger.info(
        "Translation timing start: job_id=%d segments=%d",
        translation_job_id,
        total_segments,
    )

    # --- Pass 1: run all memory checks across every segment upfront ---
    exact_memory_matches: dict[int, TranslationMemoryMatch] = {}
    semantic_suggestions: dict[int, TranslationMemoryMatch] = {}
    missing_segments: list[DocumentSegment] = []
    glossary_matches_by_segment: dict[int, list[dict]] = {}

    for s in segments:
        glossary_matches_by_segment[s.id] = _match_glossary_terms_for_segment(
            source_text=s.source_text,
            glossary_terms=glossary_candidates,
        )
        exact_match = find_exact_memory_match(
            db=db,
            source_text=s.source_text,
            source_language=source_lang,
            target_language=doc.target_language,
            customer_id=job.customer_id,
            industry=doc.industry,
            domain=doc.domain,
        )
        if exact_match:
            exact_memory_matches[s.id] = TranslationMemoryMatch(
                approved=exact_match,
                match_type="exact",
                similarity=1.0,
            )
            continue

        semantic_match = find_semantic_memory_match(
            db=db,
            source_text=s.source_text,
            source_language=source_lang,
            target_language=doc.target_language,
            customer_id=job.customer_id,
            industry=doc.industry,
            domain=doc.domain,
        )
        if semantic_match:
            semantic_suggestions[s.id] = semantic_match
        missing_segments.append(s)

    # --- Pass 2+3 merged: translate each block then commit immediately ---
    # Translating and writing block-by-block keeps progress_completed_segments
    # incrementing in real time rather than jumping 0→total after a single batch.
    segments_by_block: dict[int | None, list[DocumentSegment]] = defaultdict(list)
    for seg in segments:
        segments_by_block[seg.block_id].append(seg)

    completed_count = 0
    for block_segs in segments_by_block.values():
        # Translate only the segments in this block that need API calls.
        block_missing = [s for s in block_segs if s.id not in exact_memory_matches]
        block_provider_results: dict[int, object] = {}
        if block_missing:
            batch_ctx = [
                SegmentContext(
                    segment_id=s.id,
                    source_text=s.source_text,
                    context_before=s.context_before,
                    context_after=s.context_after,
                    glossary_terms=glossary_matches_by_segment.get(s.id, []),
                )
                for s in block_missing
            ]
            try:
                block_results_raw = provider.translate_batch(
                    segments=batch_ctx,
                    source_language=source_lang,
                    target_language=doc.target_language,
                    industry=doc.industry,
                    domain=doc.domain,
                    translation_style=translation_style,
                )
            except Exception as batch_err:
                logger.warning(
                    "Stage fallback: batch translation failed for job=%d, falling back to single segments: %s",
                    translation_job_id,
                    batch_err,
                )
                block_results_raw = [
                    provider.translate(
                        source_text=s.source_text,
                        source_language=source_lang,
                        target_language=doc.target_language,
                        industry=doc.industry,
                        domain=doc.domain,
                        context_before=s.context_before,
                        context_after=s.context_after,
                        glossary_terms=glossary_matches_by_segment.get(s.id, []),
                        translation_style=translation_style,
                    )
                    for s in block_missing
                ]
            for seg, res in zip(block_missing, block_results_raw, strict=True):
                block_provider_results[seg.id] = res

        # Build TranslationResult rows for this block and commit immediately so
        # the frontend can see progress before the full job finishes.
        block_results: list[TranslationResult] = []
        for seg in block_segs:
            if seg.id in exact_memory_matches:
                match = exact_memory_matches[seg.id]
                approved = match.approved
                block_results.append(
                    TranslationResult(
                        job_id=translation_job_id,
                        segment_id=seg.id,
                        primary_translation=approved.approved_translation,
                        final_translation=approved.approved_translation,
                        confidence_score=match.similarity,
                        # Memory can prefill a suggestion, but review must remain job-isolated and unresolved.
                        review_status="unreviewed",
                        exact_memory_used=True,
                        semantic_memory_used=False,
                        semantic_memory_details=None,
                        ambiguity_detected=False,
                        ambiguity_details=None,
                        glossary_applied=False,
                        glossary_matches=None,
                    )
                )
            else:
                res = block_provider_results[seg.id]
                seg_glossary_matches = glossary_matches_by_segment.get(seg.id, [])
                semantic_suggestion = semantic_suggestions.get(seg.id)
                semantic_details = (
                    {
                        "match_type": "semantic_memory",
                        "suggested_translation": semantic_suggestion.approved.approved_translation,
                        "similarity_score": semantic_suggestion.similarity,
                        "source_text": semantic_suggestion.approved.source_text,
                    }
                    if semantic_suggestion
                    else None
                )
                block_results.append(
                    TranslationResult(
                        job_id=translation_job_id,
                        segment_id=seg.id,
                        primary_translation=res.primary_translation,
                        final_translation=res.primary_translation,
                        confidence_score=None,
                        review_status="unreviewed",
                        exact_memory_used=False,
                        semantic_memory_used=bool(semantic_suggestion),
                        semantic_memory_details=semantic_details,
                        ambiguity_detected=res.ambiguity_detected,
                        ambiguity_details=res.ambiguity_details,
                        glossary_applied=bool(seg_glossary_matches),
                        glossary_matches=_build_glossary_matches_payload(seg_glossary_matches),
                    )
                )
        db.add_all(block_results)
        completed_count += len(block_segs)
        job.progress_completed_segments = completed_count
        # Accumulate token counts from provider results for cost tracking
        for res in block_provider_results.values():
            job.token_count_input = (job.token_count_input or 0) + res.input_tokens
            job.token_count_output = (job.token_count_output or 0) + res.output_tokens
        db.commit()

    translate_end_time = datetime.utcnow()
    translate_duration = (translate_end_time - start_time).total_seconds()
    segments_per_second = len(missing_segments) / translate_duration if translate_duration > 0 else 0
    logger.info(
        "Translation timing complete: job_id=%d segments=%d translated=%d duration=%.1fs rate=%.1f seg/s",
        translation_job_id,
        total_segments,
        len(missing_segments),
        translate_duration,
        segments_per_second,
    )

    # Final metadata update after all blocks are committed.
    job.translation_provider = provider_name
    job.translation_batch_size = len(missing_segments)
    job.status = JOB_STATUS_TRANSLATING
    job.progress_total_segments = total_segments
    job.progress_completed_segments = total_segments
    job.error_message = None
    # Calculate estimated API cost from accumulated token counts
    from pricing import calculate_api_cost
    if job.token_count_input or job.token_count_output:
        job.estimated_api_cost_usd = calculate_api_cost(
            model=provider.name if hasattr(provider, '_model') else "claude-sonnet-4-20250514",
            input_tokens=job.token_count_input or 0,
            output_tokens=job.token_count_output or 0,
        )
    db.commit()

    total_end_time = datetime.utcnow()
    total_duration = (total_end_time - start_time).total_seconds()
    logger.info(
        "Translation timing total: job_id=%d total_duration=%.1fs",
        translation_job_id,
        total_duration,
    )


def _execute_ambiguity_stage(db: Session, translation_job_id: int):
    job = db.query(TranslationJob).filter(TranslationJob.id == translation_job_id).first()
    if not job:
        raise ValueError("Translation job not found")

    results = db.query(TranslationResult).filter(TranslationResult.job_id == translation_job_id).all()
    segments_by_id = {
        segment.id: segment
        for segment in db.query(DocumentSegment).filter(DocumentSegment.document_id == job.document_id).all()
    }
    for result in results:
        segment = segments_by_id.get(result.segment_id)
        if not segment:
            continue
        _replace_segment_annotations(db, segment, result)
    db.commit()


def _execute_reconstruction_stage(db: Session, translation_job_id: int):
    job = db.query(TranslationJob).filter(TranslationJob.id == translation_job_id).first()
    if not job:
        raise ValueError("Translation job not found")
    doc = db.query(Document).filter(Document.id == job.document_id).first()
    if not doc:
        raise ValueError("Document not found")

    touched_block_ids: set[int] = set()
    segments = db.query(DocumentSegment).filter(DocumentSegment.document_id == doc.id).all()
    for segment in segments:
        if segment.block_id is not None:
            touched_block_ids.add(segment.block_id)
    for block_id in touched_block_ids:
        _refresh_document_block_translation(db, block_id, translation_job_id)

    job.status = JOB_STATUS_IN_REVIEW
    job.error_message = None
    # Increment org lifetime stats and write job_completed event
    _segs = db.query(DocumentSegment).filter(DocumentSegment.document_id == job.document_id).all()
    word_count = sum(len((s.source_text or "").split()) for s in _segs)
    if job.org_id:
        org = db.query(Organisation).filter(Organisation.id == job.org_id).first()
        if org:
            org.words_translated_lifetime = (org.words_translated_lifetime or 0) + word_count
            org.jobs_completed_lifetime = (org.jobs_completed_lifetime or 0) + 1
    from models import JobEvent
    db.add(JobEvent(
        job_id=job.id,
        event_type="job_completed",
        message=f"Translation completed with {word_count} words",
        metadata_json={"word_count": word_count},
    ))
    db.commit()

    # Auto-extract glossary suggestions in the background (fire and forget)
    try:
        from routers.glossary_suggestions import _extract_suggestions_for_job
        _extract_suggestions_for_job(db, job, job.org_id or 0)
    except Exception:
        logger.exception("Auto glossary extraction failed for job_id=%d (non-blocking)", job.id)


@router.post("/documents/{document_id}/translation-jobs", response_model=TranslationJobResponse)
def create_translation_job(
    document_id: int,
    payload: TranslationJobCreateRequest | None = None,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_active_user),
    current_org: Organisation = Depends(get_current_org),
):
    """Create a translation job and queue staged background processing."""
    doc = db.query(Document).filter(Document.id == document_id, Document.org_id == current_org.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if doc.status != "parsed":
        raise HTTPException(
            status_code=400,
            detail="Document must be parsed before creating a translation job",
        )

    segments = (
        db.query(DocumentSegment)
        .filter(DocumentSegment.document_id == document_id)
        .order_by(DocumentSegment.segment_index)
        .all()
    )
    if not segments:
        raise HTTPException(
            status_code=400,
            detail="Document has no parsed segments",
        )

    source_lang = doc.source_language or "unknown"
    translation_style = _normalize_translation_style(payload.translation_style if payload else None)
    job = TranslationJob(
        document_id=document_id,
        source_language=source_lang,
        target_language=doc.target_language,
        customer_id=doc.customer_id,
        industry=payload.industry if payload and payload.industry else doc.industry,
        domain=payload.domain if payload and payload.domain else doc.domain,
        translation_style=translation_style,
        formality=payload.formality if payload else "neutral",
        review_mode=payload.review_mode if payload else "autopilot",
        preserve_formatting=payload.preserve_formatting if payload else True,
        glossary_enabled=payload.glossary_enabled if payload else True,
        status=JOB_STATUS_TRANSLATION_QUEUED,
        translation_provider=None,
        translation_batch_size=None,
        error_message=None,
        last_saved_at=None,
        progress_total_segments=None,
        progress_completed_segments=0,
        progress_started_at=None,
        org_id=current_org.id if current_org is not None else None,
    )
    db.add(job)
    db.flush()
    _queue_stage_job(db, document_id=document_id, translation_job_id=job.id, stage_name=TRANSLATION_STAGE)
    _queue_stage_job(db, document_id=document_id, translation_job_id=job.id, stage_name=AMBIGUITY_STAGE)
    _queue_stage_job(db, document_id=document_id, translation_job_id=job.id, stage_name=RECONSTRUCTION_STAGE)
    job.progress_total_segments = len(segments)
    job.progress_completed_segments = 0
    db.commit()
    db.refresh(job)
    uid = current_user.id if current_user is not None else None
    record_event(db, JOB_CREATED, user_id=uid, job_id=job.id, document_id=document_id, org_id=job.org_id)
    run_translation_pipeline.delay(job.id, uid, current_org.id if current_org is not None else None)
    return job


@router.get("/{job_id}", response_model=TranslationJobResponse)
def get_translation_job(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    """Get a translation job by id."""
    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    doc = db.query(Document).filter(Document.id == job.document_id).first()
    response = TranslationJobResponse.model_validate(job)
    return response.model_copy(update={"document_name": doc.filename if doc else None})


@router.get("/{job_id}/progress", response_model=TranslationProgressResponse)
def get_translation_progress(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    return _calculate_translation_progress(db, job)


@router.get("/{job_id}/results", response_model=list[TranslationResultResponse])
def list_translation_results(
    job_id: int,
    filter: str = Query(default="all"),
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    """List translation results for a job."""
    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    if filter not in {"all", "ambiguities", "semantic-memory", "glossary", "memory", "flagged"}:
        raise HTTPException(status_code=400, detail="Invalid filter value")
    results = (
        db.query(TranslationResult)
        .join(DocumentSegment, TranslationResult.segment_id == DocumentSegment.id)
        .filter(TranslationResult.job_id == job_id)
        .order_by(DocumentSegment.segment_index)
        .all()
    )
    # Load segment for each result to include source text
    out = []
    for r in results:
        if filter == "ambiguities" and not getattr(r, "ambiguity_detected", False):
            continue
        if filter == "ambiguities" and _is_acceptable_final_status(getattr(r, "review_status", "")):
            continue
        if filter == "semantic-memory" and not _is_semantic_memory_result(r):
            continue
        if filter == "semantic-memory" and _is_acceptable_final_status(getattr(r, "review_status", "")):
            continue
        if filter == "glossary" and not getattr(r, "glossary_applied", False):
            continue
        if filter == "glossary" and _is_acceptable_final_status(getattr(r, "review_status", "")):
            continue
        if filter == "memory" and not _has_memory_signal(r):
            continue
        if filter == "memory" and _is_acceptable_final_status(getattr(r, "review_status", "")):
            continue
        if filter == "flagged" and not _is_flagged_result(r):
            continue
        seg = db.query(DocumentSegment).filter(DocumentSegment.id == r.segment_id).first()
        out.append(_serialize_translation_result(r, seg))
    return out


@router.get("/{job_id}/review-blocks", response_model=ReviewBlocksPageResponse)
def list_review_blocks(
    job_id: int,
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=10, ge=1, le=50),
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    """Return a page of translated document blocks with nested segment results and annotations.

    Blocks that have no TranslationResult rows yet (still translating) are returned with
    empty segments so the frontend can show placeholders while polling.
    """
    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")

    total_blocks: int = (
        db.query(func.count(DocumentBlock.id))
        .filter(DocumentBlock.document_id == job.document_id)
        .scalar()
    ) or 0
    total_pages = max(1, math.ceil(total_blocks / page_size))
    offset = (page - 1) * page_size

    blocks = (
        db.query(DocumentBlock)
        .filter(DocumentBlock.document_id == job.document_id)
        .order_by(DocumentBlock.block_index)
        .offset(offset)
        .limit(page_size)
        .all()
    )

    if not blocks:
        return ReviewBlocksPageResponse(
            blocks=[],
            page=page,
            page_size=page_size,
            total_blocks=total_blocks,
            total_pages=total_pages,
            job_status=job.status,
        )

    block_ids = {block.id for block in blocks}

    segments = (
        db.query(DocumentSegment)
        .filter(
            DocumentSegment.document_id == job.document_id,
            DocumentSegment.block_id.in_(block_ids),
        )
        .order_by(DocumentSegment.segment_index)
        .all()
    )
    segment_ids = {seg.id for seg in segments}

    results = (
        db.query(TranslationResult)
        .filter(
            TranslationResult.job_id == job_id,
            TranslationResult.segment_id.in_(segment_ids),
        )
        .all()
    ) if segment_ids else []

    results_by_segment_id = {result.segment_id: result for result in results}

    annotations = (
        db.query(SegmentAnnotation)
        .filter(
            SegmentAnnotation.segment_id.in_(segment_ids),
            SegmentAnnotation.translation_job_id == job_id,
        )
        .order_by(SegmentAnnotation.id)
        .all()
    ) if segment_ids else []

    annotations_by_segment_id: dict[int, list[SegmentAnnotation]] = defaultdict(list)
    for annotation in annotations:
        annotations_by_segment_id[annotation.segment_id].append(annotation)

    segments_by_block_id: dict[int | None, list[DocumentSegment]] = defaultdict(list)
    for segment in segments:
        segments_by_block_id[segment.block_id].append(segment)

    review_blocks: list[ReviewBlockResponse] = []
    for block in blocks:
        block_segments = []
        resolved_block_segments: list[DocumentSegment] = []
        for segment in segments_by_block_id.get(block.id, []):
            result = results_by_segment_id.get(segment.id)
            if not result:
                continue
            resolved_block_segments.append(segment)
            block_segments.append(
                _serialize_review_segment(
                    segment=segment,
                    result=result,
                    annotations=annotations_by_segment_id.get(segment.id, []),
                )
            )
        translated_text_raw, translated_text_display, translated_text_final = _build_review_block_translated_representations(
            block=block,
            block_segments=resolved_block_segments,
            result_by_segment_id=results_by_segment_id,
        )
        review_blocks.append(
            ReviewBlockResponse(
                id=block.id,
                document_id=block.document_id,
                block_index=block.block_index,
                block_type=block.block_type,
                source_text_raw=block.text_original,
                source_text_display=_clean_review_display_text(block.text_original),
                translated_text_raw=translated_text_raw,
                translated_text_display=translated_text_display,
                text_original=block.text_original,
                text_translated=translated_text_final,
                formatting_json=block.formatting_json,
                source_edited=block.source_edited,
                segments=block_segments,
            )
        )

    return ReviewBlocksPageResponse(
        blocks=review_blocks,
        page=page,
        page_size=page_size,
        total_blocks=total_blocks,
        total_pages=total_pages,
        job_status=job.status,
    )


@router.get("/{job_id}/review-summary", response_model=ReviewSummaryResponse)
def get_review_summary(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    return _calculate_review_summary(db, job)


@router.post("/{job_id}/save-draft", response_model=ReviewSummaryResponse)
def save_review_draft(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    if job.status == JOB_STATUS_EXPORTED:
        raise HTTPException(status_code=400, detail="Cannot save draft after export")
    summary = _calculate_review_summary(db, job)
    job.status = JOB_STATUS_REVIEW_COMPLETE if summary.review_complete else JOB_STATUS_DRAFT_SAVED
    job.last_saved_at = datetime.utcnow()
    db.commit()
    db.refresh(job)
    return _calculate_review_summary(db, job)


@router.post("/{job_id}/approve-safe-segments", response_model=ReviewSummaryResponse)
def approve_safe_segments(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    if job.status == JOB_STATUS_EXPORTED:
        raise HTTPException(status_code=400, detail="Cannot approve segments after export")

    results = db.query(TranslationResult).filter(TranslationResult.job_id == job.id).all()
    changed = 0
    for result in results:
        if _is_acceptable_final_status(result.review_status):
            continue
        if bool(result.ambiguity_detected):
            continue
        if bool(result.semantic_memory_used):
            continue
        if not (result.final_translation or "").strip():
            continue
        result.review_status = "approved"
        changed += 1

    if changed > 0 and job.status not in {JOB_STATUS_READY_FOR_EXPORT, JOB_STATUS_EXPORTED}:
        summary = _calculate_review_summary(db, job)
        job.status = JOB_STATUS_REVIEW_COMPLETE if summary.review_complete else JOB_STATUS_IN_REVIEW

    db.commit()
    logger.info("Bulk approved safe segments for translation_job_id=%d count=%d", job.id, changed)
    return _calculate_review_summary(db, job)


@router.post("/{job_id}/approve-all-segments", response_model=ReviewSummaryResponse)
def approve_all_segments(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    """Force-approve all segments regardless of ambiguity or memory status.

    Used by Autopilot 'Download anyway' when the user accepts the translation
    despite unresolved issues.
    """
    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    if job.status == JOB_STATUS_EXPORTED:
        raise HTTPException(status_code=400, detail="Cannot approve segments after export")

    results = db.query(TranslationResult).filter(TranslationResult.job_id == job.id).all()
    changed = 0
    for result in results:
        if _is_acceptable_final_status(result.review_status):
            continue
        if not (result.final_translation or "").strip():
            continue
        result.review_status = "approved"
        changed += 1

    if changed > 0 and job.status not in {JOB_STATUS_READY_FOR_EXPORT, JOB_STATUS_EXPORTED}:
        summary = _calculate_review_summary(db, job)
        job.status = JOB_STATUS_REVIEW_COMPLETE if summary.review_complete else JOB_STATUS_IN_REVIEW

    db.commit()
    logger.info("Force-approved all segments for translation_job_id=%d count=%d", job.id, changed)
    return _calculate_review_summary(db, job)


@router.post("/{job_id}/ready-for-export", response_model=ReviewSummaryResponse)
def mark_ready_for_export(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    if job.status == JOB_STATUS_EXPORTED:
        raise HTTPException(status_code=400, detail="Job is already exported")

    summary = _calculate_review_summary(db, job)
    if not summary.review_complete:
        raise HTTPException(
            status_code=400,
            detail="All required review items must be resolved before marking ready for export",
        )

    job.status = JOB_STATUS_READY_FOR_EXPORT
    from services.events import record_job_event
    record_job_event(db, job.id, "review_complete", "All blocks resolved, ready for export")
    db.commit()
    db.refresh(job)
    if job.org_id is not None:
        from models import OrgWebhook
        hooks = [(h.url, h.secret) for h in db.query(OrgWebhook).filter(OrgWebhook.org_id == job.org_id, OrgWebhook.is_active.is_(True)).all()]
        threading.Thread(
            target=deliver_webhook,
            args=(hooks, "job.ready_for_export", {"job_id": job.id, "document_id": job.document_id}),
            daemon=True,
        ).start()
    return _calculate_review_summary(db, job)


@router.post("/{job_id}/mark-ready", response_model=ReviewSummaryResponse)
def mark_ready_for_export_alias(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    return mark_ready_for_export(job_id=job_id, db=db, current_org=current_org)


@router.post("/{job_id}/reopen-review", response_model=ReviewSummaryResponse)
def reopen_review(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    if job.status not in {JOB_STATUS_READY_FOR_EXPORT, JOB_STATUS_EXPORTED}:
        raise HTTPException(status_code=400, detail="Only ready_for_export or exported jobs can be reopened")
    job.status = JOB_STATUS_IN_REVIEW
    db.commit()
    db.refresh(job)
    return _calculate_review_summary(db, job)


@router.post("/{job_id}/export", response_model=ExportResponse)
def export_translation_job(
    job_id: int,
    file_type: str | None = Query(default=None),
    formatting_mode: str | None = Query(default=None),
    export_format: str | None = Query(default=None),
    export_mode: str | None = Query(default=None),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_active_user),
    current_org: Organisation = Depends(get_current_org),
):
    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    raw_mode = formatting_mode if formatting_mode is not None else export_mode
    raw_format = file_type if file_type is not None else export_format
    normalized_export_mode = _normalize_export_mode(raw_mode or "preserve_formatting")
    normalized_export_format = (raw_format or "docx").strip().lower()
    if normalized_export_format not in SUPPORTED_EXPORT_FORMATS:
        raise HTTPException(status_code=400, detail="file_type must be one of: docx, rtf, txt")
    if job.status not in {JOB_STATUS_READY_FOR_EXPORT, JOB_STATUS_EXPORTED}:
        raise HTTPException(status_code=400, detail="Job must be ready_for_export or exported before export")

    doc = db.query(Document).filter(Document.id == job.document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    version = _next_export_version(job, doc)
    prefix = _export_filename_prefix(doc, job)
    filename = f"{prefix}-v{version}.{normalized_export_format}"
    filepath = EXPORT_DIR / filename
    effective_export_mode = normalized_export_mode if normalized_export_format != "txt" else "clean_text"
    if normalized_export_format == "docx":
        _build_export_docx(db, job, effective_export_mode, filepath)
    elif normalized_export_format == "rtf":
        filepath.write_text(_build_export_rtf(db, job, effective_export_mode), encoding="utf-8")
    else:
        filepath.write_text(_build_export_text(db, job, effective_export_mode), encoding="utf-8")

    exported_at = datetime.utcnow()
    _write_export_metadata(
        filename=filename,
        generated_at=exported_at,
        export_format=normalized_export_format,
        export_mode=effective_export_mode,
    )
    job.status = JOB_STATUS_EXPORTED
    job.last_saved_at = exported_at
    from services.events import record_job_event
    record_job_event(db, job.id, "export_complete", f"Exported as {normalized_export_format}", {"filename": filename, "format": normalized_export_format, "mode": effective_export_mode})
    db.commit()
    uid = current_user.id if current_user is not None else None
    record_event(db, JOB_EXPORTED, user_id=uid, job_id=job_id, document_id=job.document_id, org_id=job.org_id)
    if job.org_id is not None:
        from models import OrgWebhook
        hooks = [(h.url, h.secret) for h in db.query(OrgWebhook).filter(OrgWebhook.org_id == job.org_id, OrgWebhook.is_active.is_(True)).all()]
        threading.Thread(
            target=deliver_webhook,
            args=(hooks, "job.exported", {"job_id": job.id, "document_id": job.document_id, "filename": filename}),
            daemon=True,
        ).start()
    return ExportResponse(
        job_id=job.id,
        status=job.status,
        export_format=normalized_export_format,
        export_mode=effective_export_mode,
        filename=filename,
        download_url=f"/translation-jobs/{job.id}/exports/{filename}",
        generated_at=exported_at,
        version=version,
    )


class OverviewBlockPreview(BaseModel):
    source_text: str
    translated_text: str
    has_issue: bool


class OverviewSummary(BaseModel):
    total_blocks: int
    issue_count: int
    glossary_match_count: int
    ambiguity_count: int
    quality_score: int
    memory_reuse_count: int = 0
    word_count: int = 0


class OverviewResponse(BaseModel):
    job_id: int
    document_id: int
    status: str
    source_language: str
    target_language: str
    document_name: str
    review_mode: str
    tone: Optional[str]
    tone_applied: str = "natural"
    summary: OverviewSummary
    blocks_preview: list[OverviewBlockPreview]
    created_at: Optional[str] = None


class ReviewModeRequest(BaseModel):
    review_mode: str = Field(..., pattern="^(autopilot|manual)$")


@router.get("/{job_id}/overview", response_model=OverviewResponse)
def get_overview(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    """Return overview data for the translation overview stage."""
    job = db.query(TranslationJob).filter(
        TranslationJob.id == job_id,
        TranslationJob.org_id == current_org.id,
        TranslationJob.deleted_at.is_(None),
    ).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")

    doc = db.query(Document).filter(Document.id == job.document_id).first()
    document_name = doc.filename if doc else f"Document #{job.document_id}"

    blocks = (
        db.query(DocumentBlock)
        .filter(DocumentBlock.document_id == job.document_id)
        .order_by(DocumentBlock.block_index)
        .all()
    )
    results = db.query(TranslationResult).filter(TranslationResult.job_id == job.id).all()
    results_by_segment = {r.segment_id: r for r in results}

    total_blocks = len(blocks)
    issue_count = 0
    glossary_match_count = 0
    ambiguity_count = 0

    for r in results:
        has_ambiguity = bool(r.ambiguity_detected)
        has_glossary = bool(r.glossary_applied)
        is_resolved = _is_acceptable_final_status(r.review_status)
        if has_ambiguity and not is_resolved:
            ambiguity_count += 1
            issue_count += 1
        elif has_glossary:
            glossary_match_count += 1

    memory_reuse_count = sum(1 for r in results if bool(r.semantic_memory_used))
    quality_score = round(((total_blocks - issue_count) / total_blocks) * 100) if total_blocks > 0 else 100
    word_count = sum(len((b.text_original or "").split()) for b in blocks)
    tone_applied = job.tone or job.translation_style or "natural"

    # First 3 blocks preview
    blocks_preview = []
    for block in blocks[:3]:
        segs = db.query(DocumentSegment).filter(DocumentSegment.block_id == block.id).all()
        translated_parts = []
        block_has_issue = False
        for seg in segs:
            result = results_by_segment.get(seg.id)
            if result:
                translated_parts.append(result.final_translation or "")
                if bool(result.ambiguity_detected) and not _is_acceptable_final_status(result.review_status):
                    block_has_issue = True
            else:
                translated_parts.append("")
        blocks_preview.append(OverviewBlockPreview(
            source_text=block.text_original or "",
            translated_text=" ".join(translated_parts).strip(),
            has_issue=block_has_issue,
        ))

    return OverviewResponse(
        job_id=job.id,
        document_id=job.document_id,
        status=job.status,
        source_language=job.source_language,
        target_language=job.target_language,
        document_name=document_name,
        review_mode=job.review_mode or "autopilot",
        tone=job.tone,
        tone_applied=tone_applied,
        summary=OverviewSummary(
            total_blocks=total_blocks,
            issue_count=issue_count,
            glossary_match_count=glossary_match_count,
            ambiguity_count=ambiguity_count,
            quality_score=quality_score,
            memory_reuse_count=memory_reuse_count,
            word_count=word_count,
        ),
        blocks_preview=blocks_preview,
        created_at=job.created_at.isoformat() if job.created_at else None,
    )


@router.patch("/{job_id}/review-mode")
def update_review_mode(
    job_id: int,
    body: ReviewModeRequest,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    """Update the review mode for a translation job."""
    job = db.query(TranslationJob).filter(
        TranslationJob.id == job_id,
        TranslationJob.org_id == current_org.id,
        TranslationJob.deleted_at.is_(None),
    ).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    job.review_mode = body.review_mode
    db.commit()
    db.refresh(job)
    return {"id": job.id, "review_mode": job.review_mode}


class DueDateRequest(BaseModel):
    due_date: date | None = None


@router.patch("/{job_id}/due-date")
def update_due_date(
    job_id: int,
    body: DueDateRequest,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    """Update the due date for a translation job."""
    job = db.query(TranslationJob).filter(
        TranslationJob.id == job_id,
        TranslationJob.org_id == current_org.id,
        TranslationJob.deleted_at.is_(None),
    ).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    job.due_date = body.due_date
    # Reset reminder flags when due date changes
    job.reminder_sent_3day = False
    job.reminder_sent_1day = False
    db.commit()
    db.refresh(job)
    return {"id": job.id, "due_date": str(job.due_date) if job.due_date else None}


class ProjectAssignRequest(BaseModel):
    project_id: int | None = None


@router.patch("/{job_id}/project")
def update_project(
    job_id: int,
    body: ProjectAssignRequest,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    """Assign or remove a translation job's document from a project."""
    job = db.query(TranslationJob).filter(
        TranslationJob.id == job_id,
        TranslationJob.org_id == current_org.id,
        TranslationJob.deleted_at.is_(None),
    ).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    doc = db.query(Document).filter(Document.id == job.document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if body.project_id is not None:
        proj = db.query(Project).filter(
            Project.id == body.project_id,
            Project.org_id == current_org.id,
            Project.deleted_at.is_(None),
        ).first()
        if not proj:
            raise HTTPException(status_code=404, detail="Project not found")
    doc.project_id = body.project_id
    db.commit()
    project_name: str | None = None
    if doc.project_id is not None:
        proj = db.query(Project).filter(Project.id == doc.project_id).first()
        if proj:
            project_name = proj.name
    return {"id": job.id, "project_id": doc.project_id, "project_name": project_name}


class JobEventResponse(BaseModel):
    id: int
    job_id: int
    event_type: str
    message: Optional[str]
    metadata_json: Optional[dict]
    created_at: str

    class Config:
        from_attributes = True


@router.get("/{job_id}/events", response_model=list[JobEventResponse])
def get_job_events(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    """Return chronological event log for a translation job."""
    from models import JobEvent
    job = db.query(TranslationJob).filter(
        TranslationJob.id == job_id,
        TranslationJob.org_id == current_org.id,
        TranslationJob.deleted_at.is_(None),
    ).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    events = (
        db.query(JobEvent)
        .filter(JobEvent.job_id == job_id)
        .order_by(JobEvent.created_at.asc())
        .all()
    )
    return [
        JobEventResponse(
            id=e.id,
            job_id=e.job_id,
            event_type=e.event_type,
            message=e.message,
            metadata_json=e.metadata_json,
            created_at=e.created_at.isoformat() if e.created_at else "",
        )
        for e in events
    ]



@router.get("/{job_id}/preview", response_model=PreviewResponse)
def preview_translation_job(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")

    doc = db.query(Document).filter(Document.id == job.document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    preview_blocks = _collect_preview_blocks(db, job)
    content_raw = _render_preview_document_text(preview_blocks, "raw")
    content_display = _render_preview_document_text(preview_blocks, "display")
    return PreviewResponse(
        job_id=job.id,
        document_name=doc.filename,
        content_raw=content_raw,
        content_display=content_display,
    )


@router.get("/{job_id}/exports", response_model=list[ExportFileResponse])
def list_exports(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    doc = db.query(Document).filter(Document.id == job.document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    return _list_export_files(job, doc)


@router.get("/{job_id}/exports/{filename}")
def download_export(
    job_id: int,
    filename: str,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")

    filepath = (EXPORT_DIR / filename).resolve()
    export_root = EXPORT_DIR.resolve()
    if not str(filepath).startswith(str(export_root)):
        raise HTTPException(status_code=400, detail="Invalid export filename")
    if not filepath.exists() or not filepath.is_file():
        raise HTTPException(status_code=404, detail="Export file not found")

    suffix = filepath.suffix.lower()
    media_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if suffix == ".docx"
        else "application/rtf"
        if suffix == ".rtf"
        else "text/plain"
    )
    return FileResponse(path=str(filepath), filename=filename, media_type=media_type)


@router.patch("/translation-results/{result_id}", response_model=TranslationResultResponse)
def update_translation_result(
    result_id: int,
    body: TranslationResultUpdateRequest,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
    current_user=Depends(get_current_active_user),
):
    """Save a reviewed translation result and record approved choices only."""
    result = db.query(TranslationResult).filter(TranslationResult.id == result_id).first()
    if not result:
        raise HTTPException(status_code=404, detail="Translation result not found")

    job = db.query(TranslationJob).filter(TranslationJob.id == result.job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    if job.status == JOB_STATUS_EXPORTED:
        raise HTTPException(status_code=400, detail="Job is exported. Re-open review to edit.")

    segment = db.query(DocumentSegment).filter(DocumentSegment.id == result.segment_id).first()
    if not segment:
        raise HTTPException(status_code=404, detail="Document segment not found")

    final_translation = body.final_translation.strip()
    review_status = _normalize_review_status(body.review_status.strip())
    if not final_translation:
        raise HTTPException(status_code=400, detail="final_translation is required")
    if not review_status:
        raise HTTPException(status_code=400, detail="review_status is required")
    if review_status not in SEGMENT_REVIEW_STATES:
        raise HTTPException(
            status_code=400,
            detail=f"review_status must be one of: {', '.join(sorted(SEGMENT_REVIEW_STATES))}",
        )

    result.final_translation = final_translation
    result.review_status = review_status

    if review_status in {"approved", "edited"}:
        store_approved_translation(
            db=db,
            source_text=segment.source_text,
            approved_translation=final_translation,
            source_language=job.source_language,
            target_language=job.target_language,
            customer_id=job.customer_id,
            industry=job.industry,
            domain=job.domain,
        )

    _replace_segment_annotations(db, segment, result)
    if segment.block_id is not None:
        _refresh_document_block_translation(db, segment.block_id, result.job_id)
    if job.status not in {JOB_STATUS_READY_FOR_EXPORT, JOB_STATUS_EXPORTED}:
        summary = _calculate_review_summary(db, job)
        job.status = JOB_STATUS_REVIEW_COMPLETE if summary.review_complete else JOB_STATUS_IN_REVIEW
    db.commit()
    db.refresh(result)
    semantic_match_found, suggested_translation, similarity_score, current_translation = _semantic_choice_payload(result)
    ambiguity_choice_found, ambiguity_source_phrase, ambiguity_options = _ambiguity_choice_payload(result)

    uid = current_user.id if current_user is not None else None
    if review_status in {"approved", "edited"}:
        from services.events import record_job_event
        record_job_event(db, result.job_id, "block_approved", f"Result {result.id} {review_status}", {"result_id": result.id, "review_status": review_status})
        db.commit()
        logger.info(
            "Saved reviewed translation result_id=%d review_status=%s and stored translation memory",
            result.id,
            result.review_status,
        )
    else:
        logger.info(
            "Saved reviewed translation result_id=%d review_status=%s without storing translation memory",
            result.id,
            result.review_status,
        )

    if final_translation != result.primary_translation:
        record_event(db, TRANSLATION_EDITED, user_id=uid, job_id=result.job_id, org_id=job.org_id)
    if getattr(result, "ambiguity_detected", False) and review_status == "approved":
        record_event(db, AMBIGUITY_RESOLVED, user_id=uid, job_id=result.job_id, org_id=job.org_id)

    return TranslationResultResponse(
        id=result.id,
        job_id=result.job_id,
        segment_id=result.segment_id,
        primary_translation=result.primary_translation,
        final_translation=result.final_translation,
        confidence_score=result.confidence_score,
        review_status=_normalize_review_status(result.review_status),
        exact_memory_used=_is_exact_memory_result(result),
        semantic_memory_used=_is_semantic_memory_result(result),
        semantic_memory_details=getattr(result, "semantic_memory_details", None),
        semantic_match_found=semantic_match_found,
        suggested_translation=suggested_translation,
        similarity_score=similarity_score,
        current_translation=current_translation,
        ambiguity_choice_found=ambiguity_choice_found,
        ambiguity_source_phrase=ambiguity_source_phrase,
        ambiguity_options=ambiguity_options,
        ambiguity_detected=getattr(result, "ambiguity_detected", False) or False,
        ambiguity_details=getattr(result, "ambiguity_details", None),
        glossary_applied=getattr(result, "glossary_applied", False) or False,
        glossary_matches=_normalize_glossary_matches_payload(getattr(result, "glossary_matches", None)),
        created_at=result.created_at,
        segment=SegmentResponse.model_validate(segment),
    )


@router.get("/{job_id}/stages", response_model=list[ProcessingStageJobResponse])
def list_translation_stage_jobs(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    return (
        db.query(ProcessingStageJob)
        .filter(ProcessingStageJob.translation_job_id == job_id)
        .order_by(ProcessingStageJob.id.asc())
        .all()
    )


@router.post("/{job_id}/retry", response_model=TranslationJobResponse)
def retry_translation_job(
    job_id: int,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")

    failed_jobs = (
        db.query(ProcessingStageJob)
        .filter(
            ProcessingStageJob.translation_job_id == job_id,
            ProcessingStageJob.status == "failed",
        )
        .all()
    )
    if not failed_jobs:
        raise HTTPException(status_code=400, detail="No failed stage to retry")

    for stage_job in failed_jobs:
        if stage_job.attempt_count >= stage_job.max_attempts:
            raise HTTPException(status_code=400, detail="One or more stages exceeded retry limit")
        stage_job.status = "queued"
        stage_job.error_message = None
        stage_job.started_at = None
        stage_job.finished_at = None

    job.status = JOB_STATUS_TRANSLATION_QUEUED
    job.error_message = None
    job.progress_completed_segments = 0
    job.progress_started_at = datetime.utcnow()
    db.commit()
    db.refresh(job)
    logger.info("Retry queued for translation_job_id=%d with %d failed stages", job_id, len(failed_jobs))
    background_tasks.add_task(_run_translation_pipeline, job.id)
    return job


@router.delete("/{job_id}", status_code=204)
def delete_translation_job(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    """Delete a translation job and its output. Preserves the source document.

    Cascade deletes every FK reference: SegmentAnnotation, GlossaryTermSuggestion,
    TranslationResult, ProcessingStageJob, JobEvent, UsageEvent, then the job.
    Resets block translations to None.
    """
    from models import GlossaryTermSuggestion, JobEvent

    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")

    document_id = job.document_id

    # Delete segment annotations referencing this job
    db.query(SegmentAnnotation).filter(SegmentAnnotation.translation_job_id == job_id).delete(synchronize_session=False)
    # Delete glossary term suggestions referencing this job
    db.query(GlossaryTermSuggestion).filter(GlossaryTermSuggestion.job_id == job_id).delete(synchronize_session=False)
    # Delete translation results
    db.query(TranslationResult).filter(TranslationResult.job_id == job_id).delete(synchronize_session=False)
    # Delete processing stage jobs
    db.query(ProcessingStageJob).filter(ProcessingStageJob.translation_job_id == job_id).delete(synchronize_session=False)
    # Delete job events (operational audit — not stats data)
    db.query(JobEvent).filter(JobEvent.job_id == job_id).delete(synchronize_session=False)
    # Detach usage events from this job (preserve for lifetime stats)
    db.query(UsageEvent).filter(UsageEvent.job_id == job_id).update(
        {"job_id": None}, synchronize_session=False
    )
    # Reset block translated text
    db.query(DocumentBlock).filter(DocumentBlock.document_id == document_id).update(
        {"text_translated": None}, synchronize_session=False,
    )
    # Delete the job itself
    db.delete(job)
    db.commit()
    logger.info("Deleted translation_job_id=%d (document preserved, document_id=%d)", job_id, document_id)


@router.post("/{job_id}/retranslate", response_model=TranslationJobResponse)
def retranslate_job(
    job_id: int,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    """Clear all translation data for a job and re-trigger the translation pipeline.

    Cascade deletes every FK reference, resets job to translation_queued,
    and dispatches a new Celery task. Source Document and blocks are preserved.
    """
    from models import GlossaryTermSuggestion, JobEvent

    job = db.query(TranslationJob).filter(TranslationJob.id == job_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None)).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")

    # Clear all FK references to this job
    db.query(SegmentAnnotation).filter(SegmentAnnotation.translation_job_id == job_id).delete(synchronize_session=False)
    db.query(GlossaryTermSuggestion).filter(GlossaryTermSuggestion.job_id == job_id).delete(synchronize_session=False)
    db.query(TranslationResult).filter(TranslationResult.job_id == job_id).delete(synchronize_session=False)
    db.query(ProcessingStageJob).filter(ProcessingStageJob.translation_job_id == job_id).delete(synchronize_session=False)
    db.query(JobEvent).filter(JobEvent.job_id == job_id).delete(synchronize_session=False)
    # Detach usage events (preserve for lifetime stats — re-translation logs new events)
    db.query(UsageEvent).filter(UsageEvent.job_id == job_id).update(
        {"job_id": None}, synchronize_session=False
    )

    # Reset block translated text
    db.query(DocumentBlock).filter(DocumentBlock.document_id == job.document_id).update(
        {"text_translated": None}, synchronize_session=False
    )

    # Reset job state
    job.status = JOB_STATUS_TRANSLATION_QUEUED
    job.error_message = None
    job.progress_completed_segments = 0
    job.progress_total_segments = None
    job.progress_started_at = datetime.utcnow()
    job.last_saved_at = None

    # Re-create pipeline stage jobs
    for stage_name in [TRANSLATION_STAGE, AMBIGUITY_STAGE, RECONSTRUCTION_STAGE]:
        db.add(ProcessingStageJob(
            document_id=job.document_id,
            translation_job_id=job.id,
            stage_name=stage_name,
            status="queued",
        ))

    db.commit()
    db.refresh(job)
    logger.info("Re-translate queued for translation_job_id=%d", job_id)
    background_tasks.add_task(_run_translation_pipeline, job.id)
    return job


# ── Glossary suggestion endpoints (job-scoped) ──────────────────────────────

@router.get("/{job_id}/suggestions")
def get_job_suggestions(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    """Return all pending glossary suggestions for this job."""
    from models import GlossaryTermSuggestion
    job = db.query(TranslationJob).filter(
        TranslationJob.id == job_id,
        TranslationJob.org_id == current_org.id,
        TranslationJob.deleted_at.is_(None),
    ).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    suggestions = (
        db.query(GlossaryTermSuggestion)
        .filter(
            GlossaryTermSuggestion.job_id == job_id,
            GlossaryTermSuggestion.status == "pending",
        )
        .all()
    )
    return [
        {
            "id": s.id, "job_id": s.job_id, "source_term": s.source_term,
            "target_term": s.target_term, "source_language": s.source_language,
            "target_language": s.target_language, "frequency": s.frequency, "status": s.status,
        }
        for s in suggestions
    ]


@router.post("/{job_id}/extract-suggestions")
def extract_suggestions(
    job_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    """Manually trigger glossary term extraction for a completed job."""
    job = db.query(TranslationJob).filter(
        TranslationJob.id == job_id,
        TranslationJob.org_id == current_org.id,
        TranslationJob.deleted_at.is_(None),
    ).first()
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")
    from routers.glossary_suggestions import _extract_suggestions_for_job
    suggestions = _extract_suggestions_for_job(db, job, current_org.id)
    return [
        {
            "id": s.id, "job_id": s.job_id, "source_term": s.source_term,
            "target_term": s.target_term, "source_language": s.source_language,
            "target_language": s.target_language, "frequency": s.frequency, "status": s.status,
        }
        for s in suggestions
    ]


@router.get("/documents/{document_id}/translation-jobs", response_model=list[TranslationJobResponse])
def list_document_translation_jobs(
    document_id: int,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    """List translation jobs for a document."""
    doc = db.query(Document).filter(Document.id == document_id, Document.org_id == current_org.id, Document.deleted_at.is_(None)).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    jobs = (
        db.query(TranslationJob)
        .filter(TranslationJob.document_id == document_id, TranslationJob.org_id == current_org.id, TranslationJob.deleted_at.is_(None))
        .order_by(TranslationJob.created_at.desc())
        .all()
    )
    return jobs


@router.get("")
@router.get("/")
def list_translation_jobs(
    limit: int = Query(default=10, ge=1, le=200),
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=10, ge=1, le=200),
    order: str = Query(default="desc", pattern="^(asc|desc)$"),
    project_id: int | None = Query(default=None),
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    """List translation jobs for the current organisation (paginated).

    Returns jobs ordered by created_at with total count for pagination.
    Supports both legacy `limit` param and new `page`/`page_size` params.
    """
    order_col = TranslationJob.created_at.desc() if order == "desc" else TranslationJob.created_at.asc()
    query = db.query(TranslationJob).filter(
        TranslationJob.org_id == current_org.id,
        TranslationJob.deleted_at.is_(None),
    )
    if project_id is not None:
        project_doc_ids = [
            d.id for d in db.query(Document.id).filter(
                Document.project_id == project_id,
                Document.org_id == current_org.id,
                Document.deleted_at.is_(None),
            ).all()
        ]
        query = query.filter(TranslationJob.document_id.in_(project_doc_ids))
    total: int = query.count()
    offset = (page - 1) * page_size
    effective_limit = min(limit, page_size)
    jobs = query.order_by(order_col).offset(offset).limit(effective_limit).all()
    doc_ids = [j.document_id for j in jobs]
    docs = (
        {d.id: d for d in db.query(Document).filter(Document.id.in_(doc_ids)).all()}
        if doc_ids
        else {}
    )
    # Resolve project names for documents that belong to a project
    proj_ids = {d.project_id for d in docs.values() if d.project_id is not None}
    projects = (
        {p.id: p.name for p in db.query(Project).filter(Project.id.in_(proj_ids)).all()}
        if proj_ids
        else {}
    )
    # Compute quality scores in bulk: count blocks and unresolved ambiguities per job
    job_ids = [j.id for j in jobs]
    job_doc_ids = {j.id: j.document_id for j in jobs}
    block_counts: dict[int, int] = {}
    if doc_ids:
        for doc_id, cnt in db.query(DocumentBlock.document_id, func.count(DocumentBlock.id)).filter(
            DocumentBlock.document_id.in_(doc_ids)
        ).group_by(DocumentBlock.document_id).all():
            block_counts[doc_id] = cnt

    issue_counts: dict[int, int] = defaultdict(int)
    if job_ids:
        ambiguous_results = (
            db.query(TranslationResult.job_id, TranslationResult.review_status)
            .filter(
                TranslationResult.job_id.in_(job_ids),
                TranslationResult.ambiguity_detected.is_(True),
            )
            .all()
        )
        for jr_job_id, review_status in ambiguous_results:
            if not _is_acceptable_final_status(review_status):
                issue_counts[jr_job_id] += 1

    result = []
    for job in jobs:
        doc = docs.get(job.document_id)
        doc_name = doc.filename if doc else None
        doc_project_id = doc.project_id if doc else None
        doc_project_name = projects.get(doc_project_id) if doc_project_id else None
        total = block_counts.get(job.document_id, 0)
        issues = issue_counts.get(job.id, 0)
        qs = round(((total - issues) / total) * 100) if total > 0 else None
        result.append(
            TranslationJobResponse.model_validate(job).model_copy(
                update={
                    "document_name": doc_name,
                    "project_id": doc_project_id,
                    "project_name": doc_project_name,
                    "quality_score": qs,
                }
            )
        )
    return {"jobs": result, "total": total, "page": page, "page_size": page_size}


# ---------------------------------------------------------------------------
# Source block editing
# ---------------------------------------------------------------------------


def _count_words(text: str) -> int:
    """Count words in a text string."""
    return len(text.split())


@router.patch("/{job_id}/blocks/{block_id}/source", response_model=SourceEditResponse)
def edit_block_source(
    job_id: int,
    block_id: int,
    body: SourceEditRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_org: Organisation = Depends(get_current_org),
):
    """Edit the source text of a block and trigger re-translation.

    Stores the original source text on first edit, updates the block,
    voids all translation results for the block's segments, and queues
    re-translation. Returns threshold warnings when cumulative edits
    approach or exceed the configured limit.
    """
    job = (
        db.query(TranslationJob)
        .filter(
            TranslationJob.id == job_id,
            TranslationJob.org_id == current_org.id,
            TranslationJob.deleted_at.is_(None),
        )
        .first()
    )
    if not job:
        raise HTTPException(status_code=404, detail="Translation job not found")

    block = (
        db.query(DocumentBlock)
        .filter(
            DocumentBlock.id == block_id,
            DocumentBlock.document_id == job.document_id,
        )
        .first()
    )
    if not block:
        raise HTTPException(status_code=404, detail="Block not found in this job's document")

    # Strip any RTF control codes that leaked into the source text
    raw_source = body.source_text.strip()
    new_source = re.sub(r'\\[a-z]+\d*\s?|\{|\}', '', raw_source).strip()
    if not new_source:
        raise HTTPException(status_code=400, detail="Source text cannot be empty after cleanup")
    if new_source == block.text_original:
        raise HTTPException(status_code=400, detail="Source text is unchanged")

    # Store original on first edit
    if not block.source_edited:
        block.original_source_text = block.text_original

    # Calculate word delta
    old_word_count = _count_words(block.text_original)
    new_word_count = _count_words(new_source)
    word_delta = abs(new_word_count - old_word_count)
    job.source_edit_word_delta = (job.source_edit_word_delta or 0) + word_delta

    # Update block source
    block.text_original = new_source
    block.source_edited = True

    # Update all segments in this block with new source text
    segments = (
        db.query(DocumentSegment)
        .filter(DocumentSegment.block_id == block_id)
        .order_by(DocumentSegment.segment_index)
        .all()
    )

    # For single-segment blocks, update the segment source text directly
    if len(segments) == 1:
        segments[0].source_text = new_source

    # Void translation results for all segments in this block.
    # Use "source_changed" status so the frontend can detect re-translation in progress.
    # Keep block.text_translated so the frontend can show the previous translation
    # as a greyed-out placeholder while re-translation runs.
    segment_ids = [s.id for s in segments]
    if segment_ids:
        db.query(TranslationResult).filter(
            TranslationResult.job_id == job_id,
            TranslationResult.segment_id.in_(segment_ids),
        ).update(
            {"review_status": "source_changed", "final_translation": ""},
            synchronize_session="fetch",
        )

    # Calculate threshold
    total_source_words = sum(
        _count_words(b.text_original)
        for b in db.query(DocumentBlock)
        .filter(DocumentBlock.document_id == job.document_id)
        .all()
    )
    edit_ratio = job.source_edit_word_delta / max(total_source_words, 1)
    threshold_warning = edit_ratio >= SOURCE_EDIT_WARNING
    threshold_exceeded = edit_ratio >= SOURCE_EDIT_THRESHOLD

    # Create stage jobs for re-translation pipeline
    for stage_name in [TRANSLATION_STAGE, AMBIGUITY_STAGE, RECONSTRUCTION_STAGE]:
        db.add(ProcessingStageJob(
            document_id=job.document_id,
            translation_job_id=job.id,
            stage_name=stage_name,
            status="queued",
        ))

    db.commit()
    db.refresh(block)
    db.refresh(job)

    # Queue re-translation in background
    background_tasks.add_task(_run_translation_pipeline, job.id)

    logger.info(
        "Source edited block_id=%d job_id=%d word_delta=%d edit_ratio=%.2f",
        block_id, job_id, word_delta, edit_ratio,
    )

    return SourceEditResponse(
        block=DocumentBlockResponse.model_validate(block),
        source_edit_word_delta=job.source_edit_word_delta,
        total_source_words=total_source_words,
        threshold_warning=threshold_warning,
        threshold_exceeded=threshold_exceeded,
    )
